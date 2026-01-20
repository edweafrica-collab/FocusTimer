"""Convert Markdown to DOCX with proper formatting."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re

def convert_md_to_docx(md_path, docx_path):
    # Read markdown content
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Create document
    doc = Document()
    
    # Set up styles
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    # Process content line by line
    lines = content.split('\n')
    in_code_block = False
    code_buffer = []
    in_table = False
    table_rows = []
    
    i = 0
    while i < len(lines):
        line = lines[i]
        
        # Handle code blocks
        if line.startswith('```'):
            if in_code_block:
                # End code block
                if code_buffer:
                    p = doc.add_paragraph()
                    p.style = 'Normal'
                    run = p.add_run('\n'.join(code_buffer))
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9)
                    p.paragraph_format.left_indent = Inches(0.3)
                code_buffer = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue
        
        if in_code_block:
            code_buffer.append(line)
            i += 1
            continue
        
        # Handle tables
        if line.startswith('|') and '|' in line[1:]:
            if not in_table:
                in_table = True
                table_rows = []
            
            # Skip separator rows (|---|---|)
            if re.match(r'^\|[\s\-:]+\|', line):
                i += 1
                continue
            
            # Parse table row
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            table_rows.append(cells)
            i += 1
            continue
        elif in_table:
            # End of table
            if table_rows:
                num_cols = max(len(row) for row in table_rows)
                table = doc.add_table(rows=len(table_rows), cols=num_cols)
                table.style = 'Table Grid'
                for row_idx, row_data in enumerate(table_rows):
                    for col_idx, cell_data in enumerate(row_data):
                        if col_idx < num_cols:
                            table.rows[row_idx].cells[col_idx].text = cell_data
                doc.add_paragraph()  # Space after table
            table_rows = []
            in_table = False
        
        # Handle headers
        if line.startswith('# '):
            p = doc.add_heading(line[2:].strip(), level=1)
            i += 1
            continue
        elif line.startswith('## '):
            p = doc.add_heading(line[3:].strip(), level=2)
            i += 1
            continue
        elif line.startswith('### '):
            p = doc.add_heading(line[4:].strip(), level=3)
            i += 1
            continue
        elif line.startswith('#### '):
            p = doc.add_heading(line[5:].strip(), level=4)
            i += 1
            continue
        
        # Handle horizontal rules
        if line.strip() in ['---', '***', '___']:
            p = doc.add_paragraph('─' * 50)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue
        
        # Handle blockquotes
        if line.startswith('> '):
            text = line[2:].strip()
            # Check for alert types
            if text.startswith('[!'):
                alert_match = re.match(r'\[!(\w+)\]', text)
                if alert_match:
                    alert_type = alert_match.group(1)
                    text = text[len(alert_match.group(0)):].strip()
                    p = doc.add_paragraph()
                    run = p.add_run(f"[{alert_type}] ")
                    run.bold = True
                    p.add_run(text)
                    p.paragraph_format.left_indent = Inches(0.3)
                    i += 1
                    continue
            p = doc.add_paragraph(text)
            p.paragraph_format.left_indent = Inches(0.3)
            i += 1
            continue
        
        # Handle bullet points
        if line.strip().startswith('- ') or line.strip().startswith('* '):
            indent = len(line) - len(line.lstrip())
            text = line.strip()[2:]
            p = doc.add_paragraph(text, style='List Bullet')
            if indent > 0:
                p.paragraph_format.left_indent = Inches(0.25 * (indent // 2 + 1))
            i += 1
            continue
        
        # Handle numbered lists
        num_match = re.match(r'^(\s*)(\d+)\.\s+(.+)$', line)
        if num_match:
            text = num_match.group(3)
            p = doc.add_paragraph(text, style='List Number')
            i += 1
            continue
        
        # Handle empty lines
        if line.strip() == '':
            doc.add_paragraph()
            i += 1
            continue
        
        # Regular paragraph - handle inline formatting
        p = doc.add_paragraph()
        process_inline_formatting(p, line)
        i += 1
    
    # Handle any remaining table
    if in_table and table_rows:
        num_cols = max(len(row) for row in table_rows)
        table = doc.add_table(rows=len(table_rows), cols=num_cols)
        table.style = 'Table Grid'
        for row_idx, row_data in enumerate(table_rows):
            for col_idx, cell_data in enumerate(row_data):
                if col_idx < num_cols:
                    table.rows[row_idx].cells[col_idx].text = cell_data
    
    # Save document
    doc.save(docx_path)
    print(f"Document saved to: {docx_path}")

def process_inline_formatting(paragraph, text):
    """Process bold, italic, and code formatting in text."""
    # Simple approach: just add text with basic formatting detection
    # Pattern: **bold**, *italic*, `code`
    
    patterns = [
        (r'\*\*(.+?)\*\*', 'bold'),
        (r'\*(.+?)\*', 'italic'),
        (r'`(.+?)`', 'code'),
    ]
    
    # For simplicity, just add the text cleaning markdown syntax
    # Remove markdown formatting chars for now
    clean_text = text
    clean_text = re.sub(r'\*\*(.+?)\*\*', r'\1', clean_text)
    clean_text = re.sub(r'\*(.+?)\*', r'\1', clean_text)
    clean_text = re.sub(r'`(.+?)`', r'\1', clean_text)
    clean_text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', clean_text)  # Links
    
    paragraph.add_run(clean_text)

if __name__ == '__main__':
    md_path = r'C:\Users\USER\.gemini\antigravity\brain\a0d79e4d-01b4-43e8-b54e-ee5fe64f3868\focustimer_ux_ui_documentation.md'
    docx_path = r'C:\Users\USER\Desktop\OfficeLocal\FocusTimer\FocusTimer_UX_UI_Documentation.docx'
    convert_md_to_docx(md_path, docx_path)
